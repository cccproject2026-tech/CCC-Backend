"""Convert DATABASE_SCHEMA_DOCUMENTATION.md to DOCX."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_runs(paragraph, text: str) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif part.startswith("[") and "](" in part:
            match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if match:
                label, url = match.groups()
                run = paragraph.add_run(f"{label} ({url})")
        else:
            paragraph.add_run(part)


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [cell.strip() for cell in line.split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells if c)


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    i = 0
    in_code = False
    code_lines: list[str] = []
    code_lang = ""

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                code_lines = []
            else:
                in_code = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                if code_lang:
                    cap = doc.add_paragraph()
                    cap_run = cap.add_run(f"({code_lang})")
                    cap_run.italic = True
                    cap_run.font.size = Pt(9)
                    cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                code_lines = []
                code_lang = ""
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            style = f"Heading {min(level, 4)}"
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            header = parse_table_row(stripped)
            i += 1
            if i < len(lines):
                sep = parse_table_row(lines[i].strip())
                if is_separator_row(sep):
                    i += 1
            rows = [header]
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i].strip()))
                i += 1

            col_count = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=col_count)
            table.style = "Table Grid"
            table.autofit = True

            for r_idx, row in enumerate(rows):
                for c_idx in range(col_count):
                    cell_text = row[c_idx] if c_idx < len(row) else ""
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_formatted_runs(p, cell_text)
                    for run in p.runs:
                        run.font.size = Pt(9)
                    if r_idx == 0:
                        set_cell_shading(cell, "D9E2F3")
                        for run in p.runs:
                            run.bold = True
            doc.add_paragraph()
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, bullet_match.group(1))
            i += 1
            continue

        numbered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered_match:
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, numbered_match.group(1))
            i += 1
            continue

        blockquote_match = re.match(r"^>\s?(.*)$", stripped)
        if blockquote_match:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_formatted_runs(p, blockquote_match.group(1))
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)
        i += 1

    doc.save(str(docx_path))


if __name__ == "__main__":
    root = Path(__file__).resolve().parent.parent
    md_file = root / "DATABASE_SCHEMA_DOCUMENTATION.md"
    out_file = root / "DATABASE_SCHEMA_DOCUMENTATION.docx"
    if len(sys.argv) > 1:
        md_file = Path(sys.argv[1])
    if len(sys.argv) > 2:
        out_file = Path(sys.argv[2])
    convert_md_to_docx(md_file, out_file)
    print(f"Created: {out_file}")
